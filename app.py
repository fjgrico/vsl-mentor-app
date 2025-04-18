# 📦 Interfaz web Streamlit para generar guiones VSL + audio con branding

import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from elevenlabs.client import ElevenLabs
from elevenlabs import Voice, VoiceSettings
import base64
import datetime
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import requests
import json

# --- Configuración de la página ---
st.set_page_config(page_title="Generador VSL PRO - Mentor Digital Pro", page_icon="🧠")

# --- Login simple por código de acceso ---
ACCESS_CODE = "mentorx3"
with st.sidebar:
    st.image("logo.png", width=150)
    st.markdown("#### Acceso privado")
    input_code = st.text_input("Introduce el código de acceso:", type="password")
    if input_code != ACCESS_CODE:
        st.warning("🔐 Código incorrecto. Por favor, contacta con Mentor Digital.")
        st.stop()

st.title("🧠 Generador VSL PRO - Mentor Digital Pro")

# --- Formulario para captación de leads ---
st.subheader("📩 Antes de comenzar, déjanos tus datos")
nombre = st.text_input("Tu nombre")
email = st.text_input("Tu email")

if not nombre or not email:
    st.warning("⚠️ Rellena tu nombre y email antes de continuar.")
    st.stop()

# --- Guardar en Google Sheets ---
scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds_dict = st.secrets["google_service_account"]
creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
client = gspread.authorize(creds)
sheet = client.open("Leads_Generador_VSL").worksheet("Leads")
fecha = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
sheet.append_row([nombre, email, fecha])

# --- Guardar en ActiveCampaign (2 pasos) ---
AC_API_KEY = "01029a7b514cdc25538ea826866e885b8fbb9871142f5b12d8f6ab43b0e7067b27092614"
AC_LIST_ID = 13
headers = {
    "Api-Token": AC_API_KEY,
    "Content-Type": "application/json"
}

try:
    # Paso 1: crear contacto
    r1 = requests.post("https://mentordigitalpro.api-us1.com/api/3/contacts", headers=headers, json={
        "contact": {
            "email": email,
            "firstName": nombre
        }
    })
    r1.raise_for_status()
    contact_id = r1.json()["contact"]["id"]

    # Paso 2: añadirlo a la lista
    r2 = requests.post("https://mentordigitalpro.api-us1.com/api/3/contactLists", headers=headers, json={
        "contactList": {
            "list": AC_LIST_ID,
            "contact": contact_id,
            "status": 1
        }
    })
    r2.raise_for_status()

except Exception as e:
    st.warning(f"⚠️ No se pudo guardar en ActiveCampaign: {e}")

# --- Formulario para generar el guión ---
st.subheader("✍️ Escribe tu guión")
guion = st.text_area("Introduce el texto para tu VSL:", height=300)

font_title = "Playfair Display"
font_body = "Open Sans"
color_primary = "#C7A16A"

# --- Función para generar el DOCX ---
def generar_docx(texto):
    doc = Document()
    doc.add_picture("logo.png", width=Inches(2))
    doc.add_heading("VSL - Métodox3", level=1)
    p = doc.add_paragraph(texto)
    p.style.font.name = font_body
    p.style.font.size = Pt(12)
    doc_path = "VSL_Metodox3.docx"
    doc.save(doc_path)
    return doc_path

# --- Función para generar el audio ElevenLabs ---
def generar_audio(texto, nombre, email):
    client = ElevenLabs(api_key=st.secrets["ELEVEN_API_KEY"])
    audio = client.generate(
        text=texto,
        model="eleven_multilingual_v2",
        voice=Voice(
            voice_id="FGY2WhTYpPnrIDTdsKH5",
            settings=VoiceSettings(stability=0.4, similarity_boost=0.8)
        )
    )
    file_path = "narracion.mp3"
    with open(file_path, "wb") as f:
        for chunk in audio:
            f.write(chunk)
    return file_path

# --- Botón para generar ---
if st.button("🚀 Generar VSL PDF + Audio"):
    if guion.strip() == "":
        st.warning("⚠️ Por favor, introduce un texto válido.")
    else:
        docx_path = generar_docx(guion)
        audio_path = generar_audio(guion, nombre, email)

        with open(docx_path, "rb") as f:
            st.download_button("📄 Descargar guión en DOCX", f, file_name="VSL_Metodox3.docx")

        with open(audio_path, "rb") as f:
            st.download_button("🎧 Descargar audio MP3", f, file_name="narracion.mp3")

        st.audio(audio_path, format="audio/mp3")
        st.success("✅ ¡Tu VSL ha sido generado con éxito!")
